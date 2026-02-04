import os
import sys
import json
import time
import logging
import tempfile
import asyncio
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, CallbackQueryHandler, MessageHandler, filters, ContextTypes
from PIL import Image, ImageSequence
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import subprocess
import shutil

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

user_data = {}
processing_files = {}
ffmpeg_cache = None
config_file = "bot_config.json"
privacy_accepted = {}

user_data_lock = asyncio.Lock()
processing_files_lock = asyncio.Lock()

def load_config():
    if os.path.exists(config_file):
        with open(config_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_config(config):
    with open(config_file, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

config = load_config()

def find_ffmpeg_cached():
    global ffmpeg_cache
    if ffmpeg_cache and os.path.exists(ffmpeg_cache):
        return ffmpeg_cache
    
    if 'ffmpeg_path' in config and os.path.exists(config['ffmpeg_path']):
        ffmpeg_cache = config['ffmpeg_path']
        return ffmpeg_cache
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    local_ffmpeg = os.path.join(current_dir, "ffmpeg.exe")
    
    if os.path.exists(local_ffmpeg):
        config['ffmpeg_path'] = local_ffmpeg
        save_config(config)
        ffmpeg_cache = local_ffmpeg
        return local_ffmpeg
    
    ffmpeg_in_path = shutil.which('ffmpeg')
    if ffmpeg_in_path:
        config['ffmpeg_path'] = ffmpeg_in_path
        save_config(config)
        ffmpeg_cache = ffmpeg_in_path
        return ffmpeg_in_path
    
    common_paths = ['ffmpeg.exe', 'ffmpeg', r'.\ffmpeg.exe']
    for path in common_paths:
        try:
            creation_flags = subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
            result = subprocess.run(
                [path, '-version'],
                capture_output=True,
                text=True,
                creationflags=creation_flags,
                timeout=3
            )
            if result.returncode == 0:
                config['ffmpeg_path'] = path
                save_config(config)
                ffmpeg_cache = path
                return path
        except:
            continue
    
    return None

def detect_file_type(file_bytes, filename):
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.gif'):
        return 'GIF'
    elif filename_lower.endswith(('.mp4', '.mov', '.avi', '.mkv', '.webm')):
        return 'video'
    elif filename_lower.endswith(('.jpg', '.jpeg')):
        return 'jpg'
    elif filename_lower.endswith('.png'):
        return 'png'
    elif filename_lower.endswith('.webp'):
        return 'webp'
    elif filename_lower.endswith('.txt'):
        return 'txt'
    elif filename_lower.endswith(('.docx', '.doc')):
        return 'docx'
    elif filename_lower.endswith(('.html', '.htm')):
        return 'html'
    
    if len(file_bytes) >= 6:
        if file_bytes[:6] in [b'GIF87a', b'GIF89a']:
            return 'GIF'
        elif file_bytes[:8] == b'\x89PNG\r\n\x1a\n':
            return 'png'
        elif file_bytes[:2] == b'\xff\xd8':
            return 'jpg'
        elif len(file_bytes) >= 12 and file_bytes[:4] == b'RIFF' and file_bytes[8:12] == b'WEBP':
            return 'webp'
    
    return 'unknown'

async def update_progress(user_id, file_index, total_files, progress, status_msg=None):
    async with processing_files_lock:
        if user_id in processing_files:
            if abs(processing_files[user_id]['progress'] - progress) < 5 and progress != 100:
                return
            processing_files[user_id]['progress'] = progress
            processing_files[user_id]['current_file'] = file_index
            processing_files[user_id]['total_files'] = total_files
    
    if status_msg:
        progress_bar = "🟩" * int(progress / 20) + "⬜" * (5 - int(progress / 20))
        text = f"🔄 **Обработка файла {file_index}/{total_files}**\n\n{progress_bar} {progress}%\n\n⏳ Пожалуйста, подождите..."
        try:
            await status_msg.edit_text(text, parse_mode='Markdown')
        except:
            pass

async def show_progress_bar(message, current, total, text=""):
    try:
        progress = int((current / total) * 100) if total > 0 else 0
        progress_bar = "🟩" * int(progress / 20) + "⬜" * (5 - int(progress / 20))
        await message.edit_text(
            f"🔄 **{text}**\n\n{progress_bar} {progress}%\n\n📊 Прогресс: {current}/{total} файлов",
            parse_mode='Markdown'
        )
    except:
        pass

async def show_main_menu_after_conversion(chat_id):
    keyboard = [
        [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
        [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
        [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')],
        [InlineKeyboardButton("❓ Помощь", callback_data='help')]
    ]
    await application.bot.send_message(
        chat_id=chat_id,
        text="🔄 **Конвертер файлов**\n\nВыберите категорию для конвертации:",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    if user_id not in privacy_accepted:
        keyboard = [[InlineKeyboardButton("✅ Я согласен", callback_data='accept_privacy')]]
        await update.message.reply_text(
            "📋 **Политика конфиденциальности**\n\nИспользуя этого бота, вы соглашаетесь с:\n• Файлы хранятся только во время конвертации\n• Содержимое не анализируется\n• Данные не передаются третьим лицам\n\nНажмите кнопку для продолжения:",
            parse_mode='Markdown',
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return
    
    keyboard = [
        [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
        [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
        [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')],
        [InlineKeyboardButton("❓ Помощь", callback_data='help')]
    ]
    await update.message.reply_text(
        "🔄 **Конвертер файлов**\n\nВыберите категорию для конвертации:",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
        [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
        [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='back_to_menu')]
    ]
    
    message = "📋 **Доступные категории:**\n\n📸 **Изображения:**\n• JPG/JPEG ↔ PNG ↔ WebP ↔ GIF\n• Максимальный размер: 20 МБ\n• До 5 файлов за раз\n• Для GIF используется только первый кадр\n\n📄 **Документы:**\n• TXT ↔ DOCX\n• HTML → TXT/DOCX\n• Максимальный размер: 10 МБ\n• До 3 файлов за раз\n\n🎬 **Видео/Аудио:**\n• GIF ↔ MP4\n• Видео → MP3/WAV/FLAC\n• Максимальный размер: 50 МБ\n• 1 файл за раз\n• Максимальная длительность GIF: 30 секунд\n\n⚠️ **Важно:**\n• Бот не хранит файлы дольше времени конвертации\n• Мы не анализируем содержимое файлов\n• Для видео требуется FFmpeg\n\n🔄 **Как пользоваться:**\n1. Выберите формат конвертации\n2. Отправьте файл(ы)\n3. Отправьте команду /convert или нажмите кнопку\n4. Получите результат\n5. /cancel для отмены"
    
    await update.message.reply_text(
        message,
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def convert_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    async with user_data_lock:
        if user_id not in user_data:
            await update.message.reply_text(
                "❌ Сначала выберите тип конвертации через меню и отправьте файлы.",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
                    [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
                    [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')]
                ])
            )
            return
        
        user_info = user_data[user_id]
    
    if len(user_info['files']) == 0:
        await update.message.reply_text("❌ Сначала отправьте файлы для конвертации.")
        return
    
    await start_conversion(update, user_info, user_id)

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    async with user_data_lock:
        if user_id in user_data:
            del user_data[user_id]
    async with processing_files_lock:
        if user_id in processing_files:
            del processing_files[user_id]
    await update.message.reply_text("Операция отменена.")

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id

    if query.data == 'accept_privacy':
        privacy_accepted[user_id] = True
        await start_from_query(query)
    
    elif query.data == 'help':
        await help_command_from_query(query)
    
    elif query.data == 'back_to_menu':
        await show_main_menu(query)
    
    elif query.data == 'back_to_category':
        await back_to_category(query, user_id)
    
    elif query.data == 'category_images':
        await show_image_categories(query)
    
    elif query.data == 'category_documents':
        await show_document_formats(query)
    
    elif query.data == 'category_video':
        await show_video_categories(query)
    
    elif query.data == 'video_conversion':
        await show_video_conversion_formats(query)
    
    elif query.data == 'audio_extraction':
        await show_audio_extraction_formats(query)
    
    elif query.data == 'jpg_category':
        await show_jpg_formats(query)
    
    elif query.data == 'png_category':
        await show_png_formats(query)
    
    elif query.data == 'webp_category':
        await show_webp_formats(query)
    
    elif query.data == 'GIF_category':
        await show_GIF_formats(query)
    
    elif query.data == 'html_category':
        await show_html_formats(query)
    
    elif query.data == 'text_category':
        await show_text_category(query)
    
    elif query.data == 'save_to_cloud':
        await show_cloud_options(query)
    
    elif query.data == 'start_conversion':
        await start_conversion_from_button(query, user_id)
    
    else:
        conversion_map = {
            'jpg_to_png': ('jpg', 'png', 20, '🖼️', 5),
            'jpg_to_webp': ('jpg', 'webp', 20, '🖼️', 5),
            'jpg_to_GIF': ('jpg', 'GIF', 20, '🖼️', 5),
            'png_to_jpg': ('png', 'jpg', 20, '🖼️', 5),
            'png_to_webp': ('png', 'webp', 20, '🖼️', 5),
            'png_to_GIF': ('png', 'GIF', 20, '🖼️', 5),
            'webp_to_jpg': ('webp', 'jpg', 20, '🖼️', 5),
            'webp_to_png': ('webp', 'png', 20, '🖼️', 5),
            'webp_to_GIF': ('webp', 'GIF', 20, '🖼️', 5),
            'GIF_to_jpg': ('GIF', 'jpg', 20, '🖼️', 5),
            'GIF_to_png': ('GIF', 'png', 20, '🖼️', 5),
            'GIF_to_webp': ('GIF', 'webp', 20, '🖼️', 5),
            
            'txt_to_docx': ('txt', 'docx', 10, '📝', 3),
            'docx_to_txt': ('docx', 'txt', 10, '📝', 3),
            'html_to_txt': ('html', 'txt', 10, '🌐', 3),
            'html_to_docx': ('html', 'docx', 10, '🌐', 3),
            
            'GIF_to_mp4': ('GIF', 'mp4', 50, '🎬', 1),
            'mp4_to_GIF': ('video', 'GIF', 50, '🎬', 1),
            'video_to_mp3': ('video', 'mp3', 50, '🎵', 1),
            'video_to_wav': ('video', 'wav', 50, '🎵', 1),
            'video_to_flac': ('video', 'flac', 50, '🎵', 1),
        }
        
        if query.data in conversion_map:
            source, target, max_mb, emoji, max_files = conversion_map[query.data]
            
            if query.data in ['GIF_to_mp4', 'mp4_to_GIF', 'video_to_mp3', 'video_to_wav', 'video_to_flac']:
                ffmpeg_path = find_ffmpeg_cached()
                if not ffmpeg_path:
                    await query.edit_message_text(
                        "❌ **FFmpeg не найден**\n\nПоложите `ffmpeg.exe` в папку с ботом или установите через Chocolatey:\n`choco install ffmpeg -y`",
                        parse_mode='Markdown',
                        reply_markup=InlineKeyboardMarkup([
                            [InlineKeyboardButton("⬅️ Назад", callback_data='category_video')]
                        ])
                    )
                    return
            
            async with user_data_lock:
                user_data[user_id] = {
                    'type': query.data,
                    'source': source,
                    'target': target,
                    'max_size': max_mb * 1024 * 1024,
                    'max_files': max_files,
                    'files': [],
                    'status_message': None
                }
            
            format_names = {
                'jpg': 'JPG/JPEG изображение',
                'png': 'PNG изображение',
                'webp': 'WebP изображение',
                'GIF': 'GIF анимация',
                'txt': 'текстовый файл',
                'docx': 'Word документ',
                'html': 'HTML файл',
                'mp4': 'MP4 видео',
                'video': 'видео файл',
                'mp3': 'MP3 аудио',
                'wav': 'WAV аудио',
                'flac': 'FLAC аудио'
            }
            
            files_text = f"Максимум файлов: {max_files}" if max_files > 1 else "Только 1 файл"
            
            warning_text = ""
            if query.data in ['GIF_to_jpg', 'GIF_to_png', 'GIF_to_webp']:
                warning_text = "\n⚠️ Используется только первый кадр GIF"
            elif query.data == 'mp4_to_GIF':
                warning_text = "\n⚠️ Telegram может отправлять GIF как MP4"
            
            await query.edit_message_text(
                f"{emoji} **{source.upper()} → {target.upper()}**\n\n📤 Отправьте файл(ы) .{source}\n📏 Максимальный размер: {max_mb} МБ\n📦 {files_text}\n\n📋 Тип: {format_names.get(source, source)}\n✅ Результат: {format_names.get(target, target)}{warning_text}\n\n💡 **Инструкция:**\n1. Отправьте файлы\n2. Когда готовы, отправьте /convert\n3. Или нажмите '🚀 Начать конвертацию'\n\n❌ Отмена: /cancel",
                parse_mode='Markdown',
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("🚀 Начать конвертацию", callback_data='start_conversion')],
                    [InlineKeyboardButton("⬅️ Назад", callback_data='back_to_category')]
                ])
            )

async def start_from_query(query):
    keyboard = [
        [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
        [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
        [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')],
        [InlineKeyboardButton("❓ Помощь", callback_data='help')]
    ]
    await query.edit_message_text(
        "🔄 **Конвертер файлов**\n\nВыберите категорию для конвертации:",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def help_command_from_query(query):
    keyboard = [
        [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
        [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
        [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='back_to_menu')]
    ]
    
    message = "📋 **Доступные категории:**\n\n📸 **Изображения:**\n• JPG/JPEG ↔ PNG ↔ WebP ↔ GIF\n• Максимальный размер: 20 МБ\n• До 5 файлов за раз\n• Для GIF используется только первый кадр\n\n📄 **Документы:**\n• TXT ↔ DOCX\n• HTML → TXT/DOCX\n• Максимальный размер: 10 МБ\n• До 3 файлов за раз\n\n🎬 **Видео/Аудио:**\n• GIF ↔ MP4\n• Видео → MP3/WAV/FLAC\n• Максимальный размер: 50 МБ\n• 1 файл за раз\n• Максимальная длительность GIF: 30 секунд\n\n⚠️ **Важно:**\n• Бот не хранит файлы дольше времени конвертации\n• Мы не анализируем содержимое файлов\n• Для видео требуется FFmpeg"
    
    await query.edit_message_text(
        message,
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_main_menu(query):
    keyboard = [
        [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
        [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
        [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')],
        [InlineKeyboardButton("❓ Помощь", callback_data='help')]
    ]
    await query.edit_message_text(
        "🔄 **Конвертер файлов**\n\nВыберите категорию для конвертации:",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def back_to_category(query, user_id):
    category_map = {
        'jpg_to_png': 'category_images', 'jpg_to_webp': 'category_images', 'jpg_to_GIF': 'category_images',
        'png_to_jpg': 'category_images', 'png_to_webp': 'category_images', 'png_to_GIF': 'category_images',
        'webp_to_jpg': 'category_images', 'webp_to_png': 'category_images', 'webp_to_GIF': 'category_images',
        'GIF_to_jpg': 'category_images', 'GIF_to_png': 'category_images', 'GIF_to_webp': 'category_images',
        'txt_to_docx': 'category_documents', 'docx_to_txt': 'category_documents',
        'html_to_txt': 'category_documents', 'html_to_docx': 'category_documents',
        'GIF_to_mp4': 'category_video', 'mp4_to_GIF': 'category_video',
        'video_to_mp3': 'category_video', 'video_to_wav': 'category_video', 'video_to_flac': 'category_video'
    }
    
    async with user_data_lock:
        if user_id in user_data:
            conv_type = user_data[user_id].get('type', '')
            if conv_type in category_map:
                if category_map[conv_type] == 'category_images':
                    await show_image_categories(query)
                elif category_map[conv_type] == 'category_documents':
                    await show_document_formats(query)
                elif category_map[conv_type] == 'category_video':
                    await show_video_categories(query)
                else:
                    await show_main_menu(query)
            else:
                await show_main_menu(query)
        else:
            await show_main_menu(query)

async def show_image_categories(query):
    keyboard = [
        [InlineKeyboardButton("🖼️ JPG/JPEG файлы", callback_data='jpg_category')],
        [InlineKeyboardButton("🖼️ PNG файлы", callback_data='png_category')],
        [InlineKeyboardButton("🖼️ WebP файлы", callback_data='webp_category')],
        [InlineKeyboardButton("🖼️ GIF файлы", callback_data='GIF_category')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='back_to_menu')]
    ]
    await query.edit_message_text(
        "📸 **Категория: Изображения**\n\nВыберите исходный формат:\n• JPG/JPEG\n• PNG\n• WebP\n• GIF\n\n📏 Максимальный размер: 20 МБ\n📦 До 5 файлов за раз\n\n⚠️ Для GIF используется только первый кадр\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_document_formats(query):
    keyboard = [
        [InlineKeyboardButton("📝 Текстовые файлы", callback_data='text_category')],
        [InlineKeyboardButton("🌐 HTML файлы", callback_data='html_category')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='back_to_menu')]
    ]
    await query.edit_message_text(
        "📄 **Категория: Документы**\n\nВыберите тип документа:\n• TXT (текстовые файлы)\n• DOCX (Word документы)\n• HTML/HTM (веб-страницы)\n\n📏 Максимальный размер: 10 МБ\n📦 До 3 файлов за раз\n\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_video_categories(query):
    ffmpeg_available = find_ffmpeg_cached() is not None
    
    keyboard = [
        [InlineKeyboardButton("🎬 Конвертация видео", callback_data='video_conversion')],
        [InlineKeyboardButton("🎵 Извлечение аудио", callback_data='audio_extraction')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='back_to_menu')]
    ]
    
    message = "🎬 **Категория: Видео/Аудио**\n\nВыберите тип операции:\n• Конвертация видео (GIF ↔ MP4)\n• Извлечение аудио из видео\n\n📏 Максимальный размер: 50 МБ\n📦 Только 1 файл за раз\n📝 Максимальная длительность GIF: 30 секунд\n\n"
    
    if ffmpeg_available:
        message += "✅ FFmpeg найден"
    else:
        message += "❌ FFmpeg не найден\n🔧 Нужен для работы с видео"
    
    await query.edit_message_text(
        message,
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_video_conversion_formats(query):
    keyboard = [
        [InlineKeyboardButton("🎬 GIF → MP4", callback_data='GIF_to_mp4')],
        [InlineKeyboardButton("🎬 Видео → GIF", callback_data='mp4_to_GIF')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_video')]
    ]
    await query.edit_message_text(
        "🎬 **Конвертация видео**\n\nВыберите направление конвертации:\n• GIF → MP4 (анимация в видео)\n• Видео → GIF (видео в анимацию)\n\n⚠️ Telegram отправляет GIF как MP4\n📏 Максимальный размер: 50 МБ\n⏱️ Максимальная длительность: 30 секунд\n📦 Только 1 файл за раз",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_audio_extraction_formats(query):
    keyboard = [
        [InlineKeyboardButton("🎵 Видео → MP3", callback_data='video_to_mp3')],
        [InlineKeyboardButton("🎵 Видео → WAV", callback_data='video_to_wav')],
        [InlineKeyboardButton("🎵 Видео → FLAC", callback_data='video_to_flac')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_video')]
    ]
    await query.edit_message_text(
        "🎵 **Извлечение аудио из видео**\n\nВыберите формат аудио:\n• Видео → MP3 (хорошее сжатие)\n• Видео → WAV (без сжатия, высокое качество)\n• Видео → FLAC (без потерь)\n\n📏 Максимальный размер: 50 МБ\n📦 Только 1 файл за раз",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_text_category(query):
    keyboard = [
        [InlineKeyboardButton("📝 TXT → DOCX", callback_data='txt_to_docx')],
        [InlineKeyboardButton("📝 DOCX → TXT", callback_data='docx_to_txt')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_documents')]
    ]
    await query.edit_message_text(
        "📝 **Текстовые файлы**\n\nВыберите направление конвертации:\n• TXT → DOCX\n• DOCX → TXT\n\n📏 Максимальный размер: 10 МБ\n📦 До 3 файлов за раз\n\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_html_formats(query):
    keyboard = [
        [InlineKeyboardButton("🌐 HTML → TXT", callback_data='html_to_txt')],
        [InlineKeyboardButton("🌐 HTML → DOCX", callback_data='html_to_docx')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_documents')]
    ]
    await query.edit_message_text(
        "🌐 **HTML файлы**\n\nВыберите направление конвертации:\n• HTML → TXT\n• HTML → DOCX\n\n📋 Поддерживаются: .html, .htm\n📏 Максимальный размер: 10 МБ\n📦 До 3 файлов за раз\n\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_jpg_formats(query):
    keyboard = [
        [InlineKeyboardButton("🖼️ JPG → PNG", callback_data='jpg_to_png')],
        [InlineKeyboardButton("🖼️ JPG → WebP", callback_data='jpg_to_webp')],
        [InlineKeyboardButton("🖼️ JPG → GIF", callback_data='jpg_to_GIF')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_images')]
    ]
    await query.edit_message_text(
        "🖼️ **Исходный формат: JPG/JPEG**\n\nВыберите целевой формат:\n• JPG → PNG\n• JPG → WebP\n• JPG → GIF\n\n📏 Максимальный размер: 20 МБ\n📦 До 5 файлов за раз\n\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_png_formats(query):
    keyboard = [
        [InlineKeyboardButton("🖼️ PNG → JPG", callback_data='png_to_jpg')],
        [InlineKeyboardButton("🖼️ PNG → WebP", callback_data='png_to_webp')],
        [InlineKeyboardButton("🖼️ PNG → GIF", callback_data='png_to_GIF')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_images')]
    ]
    await query.edit_message_text(
        "🖼️ **Исходный формат: PNG**\n\nВыберите целевой формат:\n• PNG → JPG\n• PNG → WebP\n• PNG → GIF\n\n📏 Максимальный размер: 20 МБ\n📦 До 5 файлов за раз\n\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_webp_formats(query):
    keyboard = [
        [InlineKeyboardButton("🖼️ WebP → JPG", callback_data='webp_to_jpg')],
        [InlineKeyboardButton("🖼️ WebP → PNG", callback_data='webp_to_png')],
        [InlineKeyboardButton("🖼️ WebP → GIF", callback_data='webp_to_GIF')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_images')]
    ]
    await query.edit_message_text(
        "🖼️ **Исходный формат: WebP**\n\nВыберите целевой формат:\n• WebP → JPG\n• WebP → PNG\n• WebP → GIF\n\n📏 Максимальный размер: 20 МБ\n📦 До 5 файлов за раз\n\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_GIF_formats(query):
    keyboard = [
        [InlineKeyboardButton("🖼️ GIF → JPG", callback_data='GIF_to_jpg')],
        [InlineKeyboardButton("🖼️ GIF → PNG", callback_data='GIF_to_png')],
        [InlineKeyboardButton("🖼️ GIF → WebP", callback_data='GIF_to_webp')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='category_images')]
    ]
    await query.edit_message_text(
        "🖼️ **Исходный формат: GIF**\n\nВыберите целевой формат:\n• GIF → JPG\n• GIF → PNG\n• GIF → WebP\n\n⚠️ Используется только первый кадр\n📏 Максимальный размер: 20 МБ\n📦 До 5 файлов за раз\n\n💡 Можно отправить несколько файлов сразу",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def show_cloud_options(query):
    keyboard = [
        [InlineKeyboardButton("☁️ Google Drive", callback_data='save_gdrive')],
        [InlineKeyboardButton("☁️ Яндекс.Диск", callback_data='save_yadisk')],
        [InlineKeyboardButton("⬅️ Назад", callback_data='back_to_menu')]
    ]
    await query.edit_message_text(
        "☁️ **Сохранение в облако**\n\nВыберите сервис для сохранения:\n• Google Drive\n• Яндекс.Диск\n\n⚠️ Функция в разработке\nСкоро будет доступна",
        parse_mode='Markdown',
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def start_conversion_from_button(query, user_id):
    async with user_data_lock:
        if user_id not in user_data:
            await query.answer("Сначала отправьте файлы!")
            return
        
        if len(user_data[user_id]['files']) == 0:
            await query.answer("Сначала отправьте файлы!")
            return
    
    await query.edit_message_text("🚀 Начинаю конвертацию...")
    await start_conversion_from_query(query, user_id)

async def start_conversion_from_query(query, user_id):
    async with user_data_lock:
        user_info = user_data[user_id]
    await process_conversion(user_info, user_id, query.message.chat_id, query.message.message_id)

async def start_conversion(update: Update, user_info, user_id):
    await process_conversion(user_info, user_id, update.message.chat_id, update.message.message_id)

async def convert_image(file_bytes, source_format, target_format):
    try:
        image = Image.open(io.BytesIO(file_bytes))
        
        if target_format in ['jpg', 'jpeg'] and image.mode in ['RGBA', 'P']:
            image = image.convert('RGB')
        elif target_format == 'png' and image.mode == 'P':
            image = image.convert('RGBA')
        
        output_buffer = io.BytesIO()
        
        save_params = {}
        if target_format == 'jpg':
            save_params['format'] = 'JPEG'
            save_params['quality'] = 95
        elif target_format == 'png':
            save_params['format'] = 'PNG'
            save_params['optimize'] = True
        elif target_format == 'webp':
            save_params['format'] = 'WEBP'
            save_params['quality'] = 90
        elif target_format == 'GIF':
            save_params['format'] = 'GIF'
            if source_format == 'GIF':
                if hasattr(image, 'is_animated') and image.is_animated:
                    image.seek(0)
        
        image.save(output_buffer, **save_params)
        output_buffer.seek(0)
        
        return output_buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Ошибка конвертации изображения: {e}")
        raise

async def convert_txt_to_docx(txt_content):
    try:
        doc = Document()
        doc.add_heading('Конвертированный документ', 0)
        
        paragraphs = txt_content.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
    except Exception as e:
        logger.error(f"Ошибка конвертации TXT в DOCX: {e}")
        raise

async def convert_docx_to_txt(docx_bytes):
    try:
        doc_buffer = io.BytesIO(docx_bytes)
        doc = Document(doc_buffer)
        
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return '\n'.join(text_content).encode('utf-8')
    except Exception as e:
        logger.error(f"Ошибка конвертации DOCX в TXT: {e}")
        raise

async def convert_html_to_txt(html_bytes):
    try:
        html_content = html_bytes.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text.encode('utf-8')
    except Exception as e:
        logger.error(f"Ошибка конвертации HTML в TXT: {e}")
        raise

async def convert_html_to_docx(html_bytes):
    try:
        txt_content = await convert_html_to_txt(html_bytes)
        return await convert_txt_to_docx(txt_content.decode('utf-8', errors='ignore'))
    except Exception as e:
        logger.error(f"Ошибка конвертации HTML в DOCX: {e}")
        raise

async def run_ffmpeg_command(cmd, timeout=120):
    try:
        logger.info(f"Запуск FFmpeg: {' '.join(cmd)}")
        
        creation_flags = subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            creationflags=creation_flags,
            timeout=timeout
        )
        
        if result.returncode != 0:
            error_msg = result.stderr[:500] if result.stderr else "Неизвестная ошибка"
            if error_msg:
                error_msg = error_msg.encode('utf-8', errors='ignore').decode('utf-8')
            logger.error(f"Ошибка FFmpeg: {error_msg}")
            raise Exception(f"Ошибка FFmpeg: {error_msg}")
        
        return True
    except subprocess.TimeoutExpired:
        raise Exception("Таймаут конвертации. Файл слишком большой или сложный.")
    except Exception as e:
        raise Exception(f"Ошибка выполнения FFmpeg: {str(e)}")

async def convert_GIF_to_mp4(input_path, output_path, user_id=None, status_msg=None):
    ffmpeg_path = find_ffmpeg_cached()
    if not ffmpeg_path:
        raise Exception("FFmpeg не найден")
    
    cmd = [
        ffmpeg_path,
        '-i', input_path,
        '-movflags', 'faststart',
        '-pix_fmt', 'yuv420p',
        '-vf', 'scale=trunc(iw/2)*2:trunc(ih/2)*2',
        '-c:v', 'libx264',
        '-preset', 'medium',
        '-crf', '23',
        '-y',
        output_path
    ]
    
    await run_ffmpeg_command(cmd, timeout=180)

async def convert_mp4_to_GIF(input_path, output_path, user_id=None, status_msg=None):
    ffmpeg_path = find_ffmpeg_cached()
    if not ffmpeg_path:
        raise Exception("FFmpeg не найден")
    
    try:
        probe_cmd = [
            ffmpeg_path,
            '-i', input_path,
            '-show_entries', 'format=duration',
            '-v', 'quiet',
            '-of', 'csv=p=0'
        ]
        
        creation_flags = subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
        result = subprocess.run(
            probe_cmd,
            capture_output=True,
            text=True,
            creationflags=creation_flags,
            timeout=10
        )
        
        duration = 0
        if result.returncode == 0 and result.stdout.strip():
            duration = float(result.stdout.strip())
            if duration > 30:
                raise Exception(f"Видео слишком длинное ({duration:.1f} сек). Максимум: 30 секунд.")
        
        filter_complex = '[0:v] fps=10,scale=320:-1:flags=lanczos,split [a][b];[a] palettegen=stats_mode=diff [p];[b][p] paletteuse=dither=bayer:bayer_scale=5:diff_mode=rectangle'
        
        if duration > 10:
            filter_complex = f'[0:v] trim=0:30, {filter_complex}'
        
        cmd = [
            ffmpeg_path,
            '-i', input_path,
            '-vf', filter_complex,
            '-loop', '0',
            '-y',
            output_path
        ]
        
        if user_id and status_msg:
            await update_progress(user_id, 1, 1, 55, status_msg)
        
        await run_ffmpeg_command(cmd, timeout=180)
        
        if user_id and status_msg:
            await update_progress(user_id, 1, 1, 75, status_msg)
            
    except Exception as e:
        try:
            logger.info(f"Пробуем упрощенный метод конвертации MP4 в GIF: {e}")
            simple_cmd = [
                ffmpeg_path,
                '-i', input_path,
                '-vf', 'fps=10,scale=320:-1:flags=lanczos',
                '-y',
                output_path
            ]
            await run_ffmpeg_command(simple_cmd, timeout=180)
        except Exception as simple_error:
            raise Exception(f"Не удалось конвертировать MP4 в GIF: {str(e)}. Упрощенный метод тоже не сработал: {str(simple_error)}")

async def convert_video_to_audio(input_path, output_path, audio_format, user_id=None, status_msg=None):
    ffmpeg_path = find_ffmpeg_cached()
    if not ffmpeg_path:
        raise Exception("FFmpeg не найден")
    
    cmd = [ffmpeg_path, '-i', input_path]
    
    if audio_format == 'mp3':
        cmd.extend(['-q:a', '2', '-map', 'a'])
    elif audio_format == 'wav':
        cmd.extend(['-acodec', 'pcm_s16le', '-ac', '2', '-ar', '44100'])
    elif audio_format == 'flac':
        cmd.extend(['-acodec', 'flac', '-compression_level', '5'])
    
    cmd.extend(['-y', output_path])
    
    if user_id and status_msg:
        await update_progress(user_id, 1, 1, 55, status_msg)
    
    await run_ffmpeg_command(cmd, timeout=180)
    
    if user_id and status_msg:
        await update_progress(user_id, 1, 1, 75, status_msg)

async def process_video_conversion(file_bytes, conv_type, original_name, user_id=None, status_msg=None):
    input_path = None
    output_path = None
    
    try:
        detected_type = detect_file_type(bytes(file_bytes), original_name)
        
        logger.info(f"Определен тип файла: {detected_type} для {original_name}")
        
        if conv_type == 'GIF_to_mp4' and detected_type != 'GIF':
            if detected_type == 'video':
                raise Exception(f"Файл {original_name} является видеофайлом (MP4), а не GIF.")
            else:
                raise Exception(f"Файл {original_name} не является GIF файлом.")
        
        if conv_type == 'mp4_to_GIF' and detected_type not in ['video', 'GIF', 'mp4']:
            raise Exception(f"Файл {original_name} не является видеофайлом.")
        
        if conv_type == 'GIF_to_mp4':
            input_ext = 'gif'
        elif conv_type == 'mp4_to_GIF':
            if detected_type == 'GIF':
                input_ext = 'gif'
            else:
                input_ext = 'mp4'
        elif conv_type in ['video_to_mp3', 'video_to_wav', 'video_to_flac']:
            input_ext = 'mp4'
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{input_ext}') as tmp_input:
            tmp_input.write(bytes(file_bytes))
            input_path = tmp_input.name
        
        if user_id and status_msg:
            await update_progress(user_id, 1, 1, 25, status_msg)
        
        if conv_type == 'GIF_to_mp4':
            output_ext = 'mp4'
            output_path = input_path + '.mp4'
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 35, status_msg)
            
            await convert_GIF_to_mp4(input_path, output_path, user_id, status_msg)
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 65, status_msg)
        
        elif conv_type == 'mp4_to_GIF':
            output_ext = 'gif'
            output_path = input_path + '.gif'
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 35, status_msg)
            
            if detected_type == 'GIF':
                logger.info(f"Файл уже является GIF, копируем без конвертации")
                with open(input_path, 'rb') as f_in:
                    with open(output_path, 'wb') as f_out:
                        f_out.write(f_in.read())
            else:
                await convert_mp4_to_GIF(input_path, output_path, user_id, status_msg)
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 65, status_msg)
        
        elif conv_type == 'video_to_mp3':
            output_ext = 'mp3'
            output_path = input_path + '.mp3'
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 35, status_msg)
            
            await convert_video_to_audio(input_path, output_path, 'mp3', user_id, status_msg)
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 65, status_msg)
        
        elif conv_type == 'video_to_wav':
            output_ext = 'wav'
            output_path = input_path + '.wav'
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 35, status_msg)
            
            await convert_video_to_audio(input_path, output_path, 'wav', user_id, status_msg)
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 65, status_msg)
        
        elif conv_type == 'video_to_flac':
            output_ext = 'flac'
            output_path = input_path + '.flac'
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 35, status_msg)
            
            await convert_video_to_audio(input_path, output_path, 'flac', user_id, status_msg)
            
            if user_id and status_msg:
                await update_progress(user_id, 1, 1, 65, status_msg)
        
        else:
            raise Exception(f"Неизвестный тип конвертации: {conv_type}")
        
        if user_id and status_msg:
            await update_progress(user_id, 1, 1, 80, status_msg)
        
        with open(output_path, 'rb') as f:
            converted_bytes = f.read()
        
        if len(converted_bytes) == 0:
            raise Exception("Результат конвертации пуст")
        
        if user_id and status_msg:
            await update_progress(user_id, 1, 1, 90, status_msg)
        
        if '.' in original_name:
            name_without_ext = original_name.rsplit('.', 1)[0]
        else:
            name_without_ext = original_name
        
        new_filename = f"{name_without_ext}_converted.{output_ext}"
        
        mime_types = {
            'mp4': 'video/mp4',
            'gif': 'image/gif',
            'mp3': 'audio/mpeg',
            'wav': 'audio/wav',
            'flac': 'audio/flac'
        }
        
        return {
            'bytes': converted_bytes,
            'filename': new_filename,
            'mime_type': mime_types.get(output_ext, 'application/octet-stream')
        }
        
    except Exception as e:
        logger.error(f"Ошибка при конвертации видео: {e}")
        raise
    finally:
        if input_path and os.path.exists(input_path):
            try:
                os.unlink(input_path)
            except:
                pass
        if output_path and os.path.exists(output_path):
            try:
                os.unlink(output_path)
            except:
                pass

async def process_conversion(user_info, user_id, chat_id, message_id):
    total_files = len(user_info['files'])
    
    if total_files == 0:
        return
    
    status_msg = await application.bot.send_message(
        chat_id=chat_id,
        text="🔄 Начинаю обработку файлов..."
    )
    
    async with processing_files_lock:
        processing_files[user_id] = {
            'progress': 0,
            'current_file': 1,
            'total_files': total_files
        }
    
    try:
        converted_files = []
        
        for idx, file_info in enumerate(user_info['files'], 1):
            try:
                await show_progress_bar(status_msg, idx-1, total_files, "Загрузка файлов...")
                
                file = await application.bot.get_file(file_info['file_id'])
                
                await show_progress_bar(status_msg, idx-1, total_files, "Скачивание файла...")
                
                file_bytes = await file.download_as_bytearray()
                
                if len(file_bytes) > user_info['max_size']:
                    max_mb = user_info['max_size'] // (1024 * 1024)
                    raise Exception(f"Файл слишком большой. Максимум: {max_mb} МБ")
                
                await show_progress_bar(status_msg, idx-1, total_files, "Конвертация файла...")
                
                source_ext = user_info['source']
                target_ext = user_info['target']
                conv_type = user_info['type']
                
                original_name = file_info['file_name']
                
                detected_type = detect_file_type(bytes(file_bytes), original_name)
                logger.info(f"Файл {original_name}: ожидаемый тип {source_ext}, определен как {detected_type}")
                
                if source_ext in ['jpg', 'jpeg', 'png', 'webp', 'GIF']:
                    if source_ext == 'GIF' and detected_type != 'GIF':
                        raise Exception(f"Файл {original_name} не является GIF.")
                    elif source_ext == 'jpg' and detected_type not in ['jpg', 'jpeg']:
                        raise Exception(f"Файл {original_name} не является JPG/JPEG.")
                    elif source_ext == 'png' and detected_type != 'png':
                        raise Exception(f"Файл {original_name} не является PNG.")
                    elif source_ext == 'webp' and detected_type != 'webp':
                        raise Exception(f"Файл {original_name} не является WebP.")
                    
                    converted_bytes = await convert_image(bytes(file_bytes), source_ext, target_ext)
                    
                    if '.' in original_name:
                        name_without_ext = original_name.rsplit('.', 1)[0]
                    else:
                        name_without_ext = original_name
                    
                    new_filename = f"{name_without_ext}_converted.{target_ext}"
                    
                    mime_types = {
                        'jpg': 'image/jpeg',
                        'png': 'image/png',
                        'webp': 'image/webp',
                        'GIF': 'image/gif'
                    }
                    
                    converted_files.append({
                        'bytes': converted_bytes,
                        'filename': new_filename,
                        'mime_type': mime_types.get(target_ext, f'image/{target_ext}')
                    })
                
                elif conv_type == 'txt_to_docx':
                    if detected_type != 'txt':
                        raise Exception(f"Файл {original_name} не является текстовым файлом.")
                    
                    txt_content = bytes(file_bytes).decode('utf-8', errors='ignore')
                    converted_bytes = await convert_txt_to_docx(txt_content)
                    
                    new_filename = f"{original_name.rsplit('.', 1)[0]}_converted.docx"
                    
                    converted_files.append({
                        'bytes': converted_bytes,
                        'filename': new_filename,
                        'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    })
                
                elif conv_type == 'docx_to_txt':
                    if detected_type not in ['docx', 'doc']:
                        raise Exception(f"Файл {original_name} не является Word документом.")
                    
                    converted_bytes = await convert_docx_to_txt(bytes(file_bytes))
                    
                    new_filename = f"{original_name.rsplit('.', 1)[0]}_converted.txt"
                    
                    converted_files.append({
                        'bytes': converted_bytes,
                        'filename': new_filename,
                        'mime_type': 'text/plain'
                    })
                
                elif conv_type == 'html_to_txt':
                    if detected_type not in ['html', 'htm']:
                        raise Exception(f"Файл {original_name} не является HTML файлом.")
                    
                    converted_bytes = await convert_html_to_txt(bytes(file_bytes))
                    
                    new_filename = f"{original_name.rsplit('.', 1)[0]}_converted.txt"
                    
                    converted_files.append({
                        'bytes': converted_bytes,
                        'filename': new_filename,
                        'mime_type': 'text/plain'
                    })
                
                elif conv_type == 'html_to_docx':
                    if detected_type not in ['html', 'htm']:
                        raise Exception(f"Файл {original_name} не является HTML файлом.")
                    
                    converted_bytes = await convert_html_to_docx(bytes(file_bytes))
                    
                    new_filename = f"{original_name.rsplit('.', 1)[0]}_converted.docx"
                    
                    converted_files.append({
                        'bytes': converted_bytes,
                        'filename': new_filename,
                        'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    })
                
                elif conv_type in ['GIF_to_mp4', 'mp4_to_GIF', 'video_to_mp3', 'video_to_wav', 'video_to_flac']:
                    await show_progress_bar(status_msg, idx, total_files, "Конвертация видео...")
                    
                    converted_data = await process_video_conversion(
                        bytes(file_bytes), 
                        conv_type, 
                        original_name, 
                        user_id, 
                        status_msg
                    )
                    converted_files.append(converted_data)
                
                await show_progress_bar(status_msg, idx, total_files, "Файл обработан")
                
            except Exception as e:
                logger.error(f"Ошибка обработки файла {idx}: {e}")
                try:
                    await status_msg.reply_text(f"❌ Ошибка при обработке файла {idx} ({original_name}): {str(e)[:100]}")
                except:
                    pass
        
        if converted_files:
            success_count = 0
            for converted_file in converted_files:
                try:
                    mime_type = converted_file['mime_type']
                    
                    if mime_type.startswith('image/'):
                        await application.bot.send_photo(
                            chat_id=chat_id,
                            photo=converted_file['bytes'],
                            caption=f"✅ {converted_file['filename']}"
                        )
                    elif mime_type.startswith('audio/'):
                        await application.bot.send_audio(
                            chat_id=chat_id,
                            audio=converted_file['bytes'],
                            title=converted_file['filename'],
                            filename=converted_file['filename']
                        )
                    elif mime_type.startswith('video/'):
                        await application.bot.send_video(
                            chat_id=chat_id,
                            video=converted_file['bytes'],
                            caption=f"✅ {converted_file['filename']}"
                        )
                    else:
                        await application.bot.send_document(
                            chat_id=chat_id,
                            document=converted_file['bytes'],
                            filename=converted_file['filename']
                        )
                    
                    success_count += 1
                    
                except Exception as e:
                    logger.error(f"Ошибка отправки файла: {e}")
                    try:
                        await status_msg.reply_text(f"❌ Не удалось отправить файл: {str(e)[:100]}")
                    except:
                        pass
            
            await status_msg.edit_text(
                f"✅ Конвертация завершена!\n📊 Успешно обработано: {success_count}/{total_files} файлов\n📁 Формат: {user_info['source'].upper()} → {user_info['target'].upper()}"
            )
            
            await show_main_menu_after_conversion(chat_id)
        else:
            await status_msg.edit_text("❌ Не удалось обработать файлы.")
        
    except Exception as e:
        logger.error(f"Ошибка при обработке файлов: {e}")
        try:
            await status_msg.edit_text(f"❌ Ошибка: {str(e)[:150]}")
        except:
            pass
    
    finally:
        async with processing_files_lock:
            if user_id in processing_files:
                del processing_files[user_id]
        async with user_data_lock:
            if user_id in user_data:
                del user_data[user_id]

async def handle_documents(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    async with user_data_lock:
        if user_id not in user_data:
            keyboard = [
                [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
                [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
                [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')]
            ]
            await update.message.reply_text(
                "❌ Сначала выберите тип конвертации через меню.",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return
        
        user_info = user_data[user_id]
    
    if len(user_info['files']) >= user_info['max_files']:
        await update.message.reply_text(
            f"❌ Достигнут максимум {user_info['max_files']} файлов.\nОтправьте /convert для начала конвертации."
        )
        return
    
    if update.message.document:
        document = update.message.document
        
        if document.file_size and document.file_size > user_info['max_size']:
            max_mb = user_info['max_size'] // (1024 * 1024)
            await update.message.reply_text(f"❌ Файл слишком большой. Максимум: {max_mb} МБ.")
            return
        
        file_name = document.file_name.lower() if document.file_name else "document"
        source_ext = user_info['source']
        
        allowed_extensions = {
            'jpg': ['.jpg', '.jpeg', '.jpe', '.jfif'],
            'png': ['.png'],
            'webp': ['.webp'],
            'GIF': ['.gif', '.gifv'],
            'txt': ['.txt', '.text'],
            'docx': ['.docx', '.doc'],
            'html': ['.html', '.htm', '.xhtml'],
            'video': ['.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv', '.mpg', '.mpeg', '.3gp']
        }
        
        if source_ext in allowed_extensions:
            if not any(file_name.endswith(ext) for ext in allowed_extensions[source_ext]):
                if source_ext == 'video':
                    await update.message.reply_text(
                        f"❌ Ожидается видеофайл. Поддерживаемые форматы: {', '.join(allowed_extensions[source_ext])}"
                    )
                else:
                    await update.message.reply_text(
                        f"❌ Неверный формат. Ожидается: {', '.join(allowed_extensions[source_ext])}"
                    )
                return
        
        file_info = {
            'file_id': document.file_id,
            'file_name': document.file_name or f"file_{len(user_info['files']) + 1}.{source_ext}",
            'file_size': document.file_size,
            'mime_type': document.mime_type,
            'message_id': update.message.message_id
        }
        
        async with user_data_lock:
            user_info['files'].append(file_info)
        
        remaining = user_info['max_files'] - len(user_info['files'])
        
        if remaining > 0:
            message = (
                f"✅ Файл добавлен!\n📦 Загружено: {len(user_info['files'])}/{user_info['max_files']}\n📝 Осталось мест: {remaining}\n\nОтправьте ещё файлы или нажмите кнопку для начала конвертации."
            )
        else:
            message = (
                f"✅ Файл добавлен!\n📦 Загружено: {len(user_info['files'])}/{user_info['max_files']}\n\n📊 Все файлы получены! Начинаем конвертацию..."
            )
        
        keyboard = [
            [InlineKeyboardButton("🚀 Начать конвертацию", callback_data='start_conversion')],
            [InlineKeyboardButton("❌ Отменить", callback_data='back_to_category')]
        ] if remaining > 0 else []
        
        await update.message.reply_text(
            message,
            reply_markup=InlineKeyboardMarkup(keyboard) if keyboard else None
        )

async def handle_photos(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    async with user_data_lock:
        if user_id not in user_data:
            keyboard = [
                [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
                [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
                [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')]
            ]
            await update.message.reply_text(
                "❌ Сначала выберите тип конвертации через меню.",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return
        
        user_info = user_data[user_id]
    
    if len(user_info['files']) >= user_info['max_files']:
        await update.message.reply_text(
            f"❌ Достигнут максимум {user_info['max_files']} файлов.\nОтправьте /convert для начала конвертации."
        )
        return
    
    if update.message.photo:
        photo = update.message.photo[-1]
        
        if photo.file_size and photo.file_size > user_info['max_size']:
            max_mb = user_info['max_size'] // (1024 * 1024)
            await update.message.reply_text(f"❌ Фото слишком большое. Максимум: {max_mb} МБ.")
            return
        
        file_info = {
            'file_id': photo.file_id,
            'file_name': f"photo_{len(user_info['files']) + 1}.jpg",
            'file_size': photo.file_size,
            'mime_type': 'image/jpeg',
            'message_id': update.message.message_id
        }
        
        async with user_data_lock:
            user_info['files'].append(file_info)
        
        remaining = user_info['max_files'] - len(user_info['files'])
        
        if remaining > 0:
            message = (
                f"✅ Фото добавлено!\n📦 Загружено: {len(user_info['files'])}/{user_info['max_files']}\n📸 Осталось мест: {remaining}\n\nОтправьте ещё фото или нажмите кнопку для начала конвертации."
            )
        else:
            message = (
                f"✅ Фото добавлено!\n📦 Загружено: {len(user_info['files'])}/{user_info['max_files']}\n\n📊 Все фото получены! Начинаем конвертацию..."
            )
        
        keyboard = [
            [InlineKeyboardButton("🚀 Начать конвертацию", callback_data='start_conversion')],
            [InlineKeyboardButton("❌ Отменить", callback_data='back_to_category')]
        ] if remaining > 0 else []
        
        await update.message.reply_text(
            message,
            reply_markup=InlineKeyboardMarkup(keyboard) if keyboard else None
        )

async def handle_video(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    async with user_data_lock:
        if user_id not in user_data:
            keyboard = [
                [InlineKeyboardButton("📸 Изображения", callback_data='category_images')],
                [InlineKeyboardButton("📄 Документы", callback_data='category_documents')],
                [InlineKeyboardButton("🎬 Видео/Аудио", callback_data='category_video')]
            ]
            await update.message.reply_text(
                "❌ Сначала выберите тип конвертации через меню.",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return
        
        user_info = user_data[user_id]
    
    if len(user_info['files']) >= user_info['max_files']:
        await update.message.reply_text(f"❌ Максимум {user_info['max_files']} файлов.")
        return
    
    if update.message.video:
        video = update.message.video
        
        if user_info['type'] == 'mp4_to_GIF' and video.duration and video.duration > 30:
            await update.message.reply_text(
                "❌ Видео слишком длинное для конвертации в GIF.\n"
                "Максимальная длительность: 30 секунд.\n"
                f"Текущая длительность: {video.duration} секунд."
            )
            return
        
        if video.file_size and video.file_size > user_info['max_size']:
            max_mb = user_info['max_size'] // (1024 * 1024)
            await update.message.reply_text(f"❌ Видео слишком большое. Максимум: {max_mb} МБ.")
            return
        
        file_info = {
            'file_id': video.file_id,
            'file_name': video.file_name or f"video_{len(user_info['files']) + 1}.mp4",
            'file_size': video.file_size,
            'mime_type': video.mime_type,
            'message_id': update.message.message_id
        }
        
        async with user_data_lock:
            user_info['files'].append(file_info)
        
        duration_text = f"{video.duration} сек" if video.duration else "неизвестно"
        size_text = f"{video.file_size // (1024*1024)} МБ" if video.file_size else "неизвестно"
        
        message = (
            f"✅ Видео добавлено!\n"
            f"📹 Размер: {size_text}\n"
            f"⏱️ Длительность: {duration_text}\n\n"
            f"Отправьте /convert для начала конвертации."
        )
        
        keyboard = [
            [InlineKeyboardButton("🚀 Начать конвертацию", callback_data='start_conversion')],
            [InlineKeyboardButton("❌ Отменить", callback_data='back_to_category')]
        ]
        
        await update.message.reply_text(
            message,
            reply_markup=InlineKeyboardMarkup(keyboard)
        )

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    async with user_data_lock:
        if user_id not in user_data:
            return
    
    text = update.message.text.lower().strip()
    
    text_commands = ['готово', 'конвертировать', 'start', 'go', 'convert', 'начать', 'старт']
    if text in text_commands:
        await convert_command(update, context)
    elif text in ['отмена', 'cancel', 'стоп', 'stop']:
        await cancel(update, context)
    elif text in ['помощь', 'help', 'справка']:
        await help_command(update, context)
    elif text in ['меню', 'menu', 'начать сначала']:
        await start(update, context)

def main():
    global application
    
    
    TOKEN = ""
    
    application = Application.builder().token(TOKEN).build()
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("convert", convert_command))
    application.add_handler(CommandHandler("cancel", cancel))
    application.add_handler(CallbackQueryHandler(button_handler))
    application.add_handler(MessageHandler(filters.PHOTO, handle_photos))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_documents))
    application.add_handler(MessageHandler(filters.VIDEO, handle_video))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    logger.info("___Бот запущен___")
    application.run_polling(allowed_updates=Update.ALL_TYPES, close_loop=False)
        
if __name__ == '__main__':
    main()